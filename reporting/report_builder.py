import datetime
import io
from typing import List, Dict, Any, Optional

# PDF imports
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

# Word imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def _get_timestamp(scan_time):
    if scan_time:
        return scan_time
    return datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")

def _get_risk_level(headers_result, http_result):
    # headers_result is a list of HeaderCheckResult objects
    # http_result is a dictionary (http_results.__dict__ passed from app.py)
    
    headers_ok = 0
    for h in headers_result:
        status = getattr(h, 'status', None)
        if status == 'ok':
            headers_ok += 1
            
    uses_https = http_result.get('uses_https', False)
    
    score = headers_ok + (1 if uses_https else 0)
    if score >= 5:
        return "Baixo", "O site demonstra boas práticas iniciais de segurança.", "008000" # Green
    elif score >= 3:
        return "Médio", "Existem pontos importantes de segurança ausentes ou configurados de forma fraca.", "FFA500" # Orange
    else:
        return "Alto", "O site carece de proteções fundamentais contra ataques comuns.", "FF0000" # Red

# --- TXT REPORT ---

def generate_txt_report(url, http_result, headers_result, cookies_result, scan_time=None) -> str:
    timestamp = _get_timestamp(scan_time)
    risk_lv, risk_msg, _ = _get_risk_level(headers_result, http_result)
    
    lines = [
        f"Relatório de Análise de Segurança – {url}",
        "=" * 50,
        f"Data/hora da análise: {timestamp}",
        "",
        "AVISO EDUCATIVO:",
        "Esta ferramenta é para fins educativos e não substitui testes de segurança profissionais.",
        "Use apenas em sites sob seu controle ou com autorização.",
        "",
        "1. VISÃO GERAL",
        "-" * 20,
        f"Nível de Risco: {risk_lv}",
        f"Resumo: {risk_msg}",
        "",
        "2. HTTPS E TRANSPORTE",
        "-" * 20,
        f"Usa HTTPS: {'Sim' if http_result.get('uses_https') else 'Não'}",
        f"Redireciona HTTP para HTTPS: {'Sim' if http_result.get('redirects_http_to_https') else 'Não'}",
        f"Certificado: {http_result.get('certificate_info')}",
        "Problemas encontrados:" if http_result.get('https_issues') else "",
    ]
    
    if http_result.get('https_issues'):
        for issue in http_result.get('https_issues'):
            lines.append(f" - {issue}")
            
    lines.extend([
        "",
        "3. CABEÇALHOS DE SEGURANÇA",
        "-" * 20,
        f"{'Cabeçalho':<30} | {'Status':<10} | {'Detalhes'}",
        "-" * 70
    ])
    
    for h in headers_result:
        name = getattr(h, 'name', 'N/A')
        status = getattr(h, 'status', 'missing').upper()
        details = getattr(h, 'details', 'N/A')
        lines.append(f"{name:<30} | {status:<10} | {details}")

    lines.extend([
        "",
        "4. COOKIES",
        "-" * 20,
        f"{'Nome':<25} | {'Secure':<7} | {'HttpOnly':<8} | {'Problemas'}",
        "-" * 70
    ])
    
    if not cookies_result:
        lines.append("Nenhum cookie detectado.")
    else:
        for c in cookies_result:
            name = getattr(c, 'name', 'N/A')
            sec = 'Sim' if getattr(c, 'secure', False) else 'Não'
            htt = 'Sim' if getattr(c, 'httponly', False) else 'Não'
            iss = ", ".join(getattr(c, 'issues', []))
            lines.append(f"{name:<25} | {sec:<7} | {htt:<8} | {iss}")

    lines.extend([
        "",
        "5. CONCLUSÃO E PRÓXIMOS PASSOS",
        "-" * 20,
        "1. Priorize a ativação total de HTTPS e HSTS.",
        "2. Corrija flags de cookies de sessão (HttpOnly e Secure).",
        "3. Implemente cabeçalhos como X-Frame-Options e CSP.",
        "",
        "--- Gerado por Mini SecurityHeaders Educativo ---"
    ])
    
    return "\n".join(lines)

# --- PDF REPORT ---

def generate_pdf_report(url, http_result, headers_result, cookies_result, scan_time=None) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], alignment=1, spaceAfter=20)
    header_style = ParagraphStyle('HeaderStyle', parent=styles['Heading2'], spaceBefore=15, spaceAfter=10, color=colors.HexColor('#1b263b'))
    normal_style = styles['Normal']
    warning_style = ParagraphStyle('WarningStyle', parent=styles['Italic'], color=colors.red, fontSize=9, alignment=1)

    elements = []
    
    # Title & Header
    elements.append(Paragraph(f"Relatório de Análise de Segurança", title_style))
    elements.append(Paragraph(f"URL: {url}", styles['Heading3']))
    elements.append(Paragraph(f"Data da Análise: {_get_timestamp(scan_time)}", normal_style))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph("Esta ferramenta é para fins educativos e não substitui testes de segurança profissionais.", warning_style))
    elements.append(Spacer(1, 20))
    
    # 1. Visão Geral
    risk_lv, risk_msg, risk_hex = _get_risk_level(headers_result, http_result)
    elements.append(Paragraph("1. Visão Geral", header_style))
    elements.append(Paragraph(f"<b>Nível de Risco:</b> <font color='#{risk_hex}'>{risk_lv}</font>", normal_style))
    elements.append(Paragraph(f"<b>Resumo:</b> {risk_msg}", normal_style))
    elements.append(Spacer(1, 10))
    
    # 2. HTTPS
    elements.append(Paragraph("2. HTTPS e Transporte", header_style))
    elements.append(Paragraph(f"Usa HTTPS: {'Sim' if http_result.get('uses_https') else 'Não'}", normal_style))
    elements.append(Paragraph(f"Redireciona HTTP para HTTPS: {'Sim' if http_result.get('redirects_http_to_https') else 'Não'}", normal_style))
    elements.append(Paragraph(f"Certificado: {http_result.get('certificate_info')}", normal_style))
    
    # 3. Cabeçalhos Table
    elements.append(Paragraph("3. Cabeçalhos de Segurança", header_style))
    h_data = [['Cabeçalho', 'Status', 'Detalhes']]
    for h in headers_result:
        name = getattr(h, 'name', 'N/A')
        status = getattr(h, 'status', 'missing').upper()
        details = getattr(h, 'details', 'N/A')
        h_data.append([name, status, Paragraph(details, normal_style)])
    
    t_headers = Table(h_data, colWidths=[140, 60, 280])
    t_headers.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
    ]))
    elements.append(t_headers)
    
    # 4. Cookies Table
    elements.append(Paragraph("4. Cookies", header_style))
    if not cookies_result:
        elements.append(Paragraph("Nenhum cookie detectado.", normal_style))
    else:
        c_data = [['Nome', 'Secure', 'HttpOnly', 'Problemas']]
        for c in cookies_result:
            name = getattr(c, 'name', 'N/A')
            sec = '✓' if getattr(c, 'secure', False) else '✗'
            htt = '✓' if getattr(c, 'httponly', False) else '✗'
            iss = ", ".join(getattr(c, 'issues', []))
            c_data.append([name, sec, htt, Paragraph(iss, normal_style)])
        
        t_cookies = Table(c_data, colWidths=[120, 50, 60, 250])
        t_cookies.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ]))
        elements.append(t_cookies)
    
    # 5. Conclusion
    elements.append(Paragraph("5. Conclusão", header_style))
    elements.append(Paragraph("Foque em manter HTTPS ativo, configurar HSTS e ajustar as flags dos cookies de sessão.", normal_style))
    
    doc.build(elements)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes

# --- WORD REPORT ---

def generate_docx_report(url, http_result, headers_result, cookies_result, scan_time=None) -> bytes:
    doc = Document()
    
    # Helper for alignment
    def set_center(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    title = doc.add_heading(f"Relatório de Análise de Segurança", 0)
    set_center(title)
    
    p = doc.add_paragraph()
    p.add_run(f"URL Analisada: {url}\n").bold = True
    p.add_run(f"Data: {_get_timestamp(scan_time)}").italic = True
    set_center(p)

    doc.add_paragraph("AVISO: Ferramenta educativa. Não substitui auditorias profissionais.").italic = True

    # 1. Overview
    risk_lv, risk_msg, risk_hex = _get_risk_level(headers_result, http_result)
    doc.add_heading('1. Visão Geral', level=1)
    ov = doc.add_paragraph()
    ov.add_run("Nível de Risco: ").bold = True
    rv = ov.add_run(risk_lv)
    rv.bold = True
    # Convert hex to RGB for python-docx
    r, g, b = int(risk_hex[0:2], 16), int(risk_hex[2:4], 16), int(risk_hex[4:6], 16)
    rv.font.color.rgb = RGBColor(r, g, b)
    doc.add_paragraph(risk_msg)

    # 2. HTTPS
    doc.add_heading('2. HTTPS e Transporte', level=1)
    doc.add_paragraph(f"Usa HTTPS: {'Sim' if http_result.get('uses_https') else 'Não'}")
    doc.add_paragraph(f"Transporte Seguro: {http_result.get('certificate_info')}")

    # 3. Headers
    doc.add_heading('3. Cabeçalhos de Segurança', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cabeçalho'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Detalhes'
    
    for h in headers_result:
        row_cells = table.add_row().cells
        row_cells[0].text = getattr(h, 'name', 'N/A')
        row_cells[1].text = getattr(h, 'status', 'missing').upper()
        row_cells[2].text = getattr(h, 'details', 'N/A')

    # 4. Cookies
    doc.add_heading('4. Cookies', level=1)
    if not cookies_result:
        doc.add_paragraph("Nenhum cookie detectado.")
    else:
        ctable = doc.add_table(rows=1, cols=4)
        ctable.style = 'Table Grid'
        chdr = ctable.rows[0].cells
        chdr[0].text = 'Nome'
        chdr[1].text = 'Secure'
        chdr[2].text = 'HttpOnly'
        chdr[3].text = 'Problemas'
        
        for c in cookies_result:
            row = ctable.add_row().cells
            row[0].text = getattr(c, 'name', 'N/A')
            row[1].text = 'Sim' if getattr(c, 'secure', False) else 'Não'
            row[2].text = 'Sim' if getattr(c, 'httponly', False) else 'Não'
            row[3].text = ", ".join(getattr(c, 'issues', []))

    # Conclusion
    doc.add_heading('5. Conclusão', level=1)
    doc.add_paragraph("Recomenda-se a implementação das correções sugeridas para mitigar riscos de ataques como XSS e Clickjacking.")

    buffer = io.BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()
    buffer.close()
    return content

# Keep old function for compatibility if needed, but the new ones follow requested naming
def build_markdown_report(url, results):
    # This remains for backward compatibility in app.py if it wasn't updated yet
    from reporting.report_builder_legacy import build_markdown_report as legacy_build
    return legacy_build(url, results)
